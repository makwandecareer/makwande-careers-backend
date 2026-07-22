from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import Response
from pypdf import PdfReader

from app.database import get_connection
from app.dependencies import get_current_user
from app.schemas_ai_cv_v4_1 import (
    ATSScoreRequest,
    ATSScoreResponse,
    ExportCVRequest,
    GenerateCVRequest,
    GeneratedCVResponse,
    ImproveExperienceRequest,
    ImproveSummaryRequest,
    ImportCVRequest,
    CVImportDraft,
    ImportedContact,
    RevisionAcceptRequest,
    TextImprovementResponse,
)
from app.services.openai_career_engine import CareerEngineError, OpenAICareerEngine
from app.services.cv_builder_v4_1 import (
    calculate_ats,
    export_docx,
    export_pdf,
    improve_experience_text,
    improve_summary_text,
)

router = APIRouter(prefix="/ai-cv", tags=["AI CV Builder"])

MAX_CV_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 150_000
SUPPORTED_CV_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _fallback_import_draft(
    text: str,
    user: dict,
    target_role: str,
) -> CVImportDraft:
    email_match = re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text, re.I)
    phone_match = re.search(r"(?:\+?\d[\d\s()-]{7,}\d)", text)
    linkedin_match = re.search(r"https?://(?:www\.)?linkedin\.com/in/[^\s]+", text, re.I)
    years = [int(value) for value in re.findall(r"\b((?:19|20)\d{2})\b", text)]
    span = max(years) - min(years) if len(years) >= 2 else 0

    if re.search(r"\b(?:chief|ceo|cfo|coo|executive director|vice president)\b", text, re.I):
        level = "executive"
        template = "executive"
    elif span >= 10 or re.search(r"\b(?:head of|director|senior manager)\b", text, re.I):
        level = "senior"
        template = "professional"
    elif span >= 4:
        level = "mid-career"
        template = "professional"
    elif re.search(r"\b(?:student|graduate|matric|internship)\b", text, re.I):
        level = "graduate"
        template = "graduate"
    else:
        level = "early-career"
        template = "ats-standard"

    questions_by_level = {
        "graduate": [
            "Which academic projects, practical assignments or research best demonstrate your ability?",
            "Have you completed internships, volunteer work, leadership activities or part-time work?",
            "Which software, tools, laboratory methods or technical skills can you use confidently?",
        ],
        "early-career": [
            "What results have you delivered in internships or your first roles?",
            "Which tools and job-specific skills can you demonstrate with evidence?",
            "What responsibilities have you taken ownership of independently?",
        ],
        "mid-career": [
            "Which achievements improved revenue, cost, quality, safety, delivery or customer outcomes?",
            "What projects, systems or teams have you led?",
            "What measurable scope—budget, team size, volume or geography—can you verify?",
        ],
        "senior": [
            "What strategic programmes and transformations have you led?",
            "What team sizes, budgets, regions and executive stakeholders were in your scope?",
            "Which quantified business outcomes best demonstrate your leadership impact?",
        ],
        "executive": [
            "What enterprise strategy, governance or board-level responsibilities can you verify?",
            "What P&L, workforce, portfolio or geographic scale have you led?",
            "Which transformations created measurable organisational value?",
        ],
    }

    return CVImportDraft(
        candidate_level=level,
        suggested_template=template,
        personal_details=ImportedContact(
            full_name=user.get("full_name", ""),
            email=email_match.group(0) if email_match else user.get("email", ""),
            phone=phone_match.group(0).strip() if phone_match else "",
            linkedin_url=linkedin_match.group(0) if linkedin_match else "",
        ),
        professional_title=target_role,
        professional_summary="",
        missing_details=[
            "Review the imported information and add verified achievements.",
            "Confirm employment and education dates before generating the final CV.",
        ],
        follow_up_questions=questions_by_level[level],
        facts_to_verify=["Automated extraction was limited; verify all imported facts."],
    )


def _build_import_draft(
    text: str,
    user: dict,
    target_role: str,
    job_description: str,
) -> tuple[CVImportDraft, str]:
    try:
        draft = OpenAICareerEngine().structured(
            schema=CVImportDraft,
            task=(
                "Extract the candidate's existing CV into a structured draft. "
                "Preserve only facts explicitly supported by the CV. Classify the career level, "
                "select an appropriate template, identify missing information and ask concise "
                "level-specific follow-up questions. Never invent achievements, dates or metrics."
            ),
            profile_bundle={
                "account_name": user.get("full_name"),
                "account_email": user.get("email"),
            },
            request_data={
                "extracted_cv": text,
                "target_role": target_role,
                "job_description": job_description,
            },
        )
        return draft, "openai-structured"
    except CareerEngineError:
        return _fallback_import_draft(text, user, target_role), "safe-fallback"


def _clean_extracted_text(value: str) -> str:
    value = value.replace("\x00", "")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line)[:MAX_EXTRACTED_CHARACTERS]


def _extract_pdf(content: bytes) -> tuple[str, int]:
    try:
        reader = PdfReader(BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return _clean_extracted_text(text), len(reader.pages)
    except Exception as exc:
        raise HTTPException(
            status_code=422,
            detail="The PDF could not be read. Upload a text-based PDF or DOCX file.",
        ) from exc


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text for cell in row.cells))
        return _clean_extracted_text("\n".join(blocks))
    except Exception as exc:
        raise HTTPException(
            status_code=422,
            detail="The DOCX file could not be read. Export it again and retry.",
        ) from exc


@router.post("/intake-analysis")
async def analyse_uploaded_cv(
    file: UploadFile = File(...),
    target_role: str = Form(default=""),
    job_description: str = Form(default=""),
    user=Depends(get_current_user),
):
    """Extract a CV securely in memory for the authenticated user's ATS analysis."""
    filename = Path(file.filename or "uploaded-cv").name
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_CV_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail="Upload a PDF, DOCX or TXT CV. Legacy DOC files must be saved as DOCX.",
        )

    content = await file.read(MAX_CV_UPLOAD_BYTES + 1)
    await file.close()

    if not content:
        raise HTTPException(status_code=422, detail="The uploaded CV is empty.")
    if len(content) > MAX_CV_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="The CV must be 10 MB or smaller.")

    page_count = None
    if extension == ".pdf":
        extracted_text, page_count = _extract_pdf(content)
    elif extension == ".docx":
        extracted_text = _extract_docx(content)
    else:
        try:
            extracted_text = _clean_extracted_text(content.decode("utf-8-sig"))
        except UnicodeDecodeError:
            extracted_text = _clean_extracted_text(content.decode("latin-1"))

    word_count = len(re.findall(r"\b[\w+#.-]+\b", extracted_text))
    if word_count < 30:
        raise HTTPException(
            status_code=422,
            detail=(
                "Very little readable text was found. If this is a scanned PDF, "
                "export it with selectable text or upload a DOCX version."
            ),
        )

    draft, generation_mode = _build_import_draft(
        extracted_text,
        user,
        target_role.strip()[:200],
        job_description.strip()[:20_000],
    )

    return {
        "filename": filename,
        "size": len(content),
        "content_type": file.content_type or "application/octet-stream",
        "text": extracted_text,
        "word_count": word_count,
        "page_count": page_count,
        "target_role": target_role.strip()[:200],
        "job_description_received": bool(job_description.strip()),
        "stored": False,
        "draft": draft.model_dump(),
        "generation_mode": generation_mode,
    }


@router.post("/import-draft", response_model=GeneratedCVResponse)
def save_imported_cv_draft(
    payload: ImportCVRequest,
    user=Depends(get_current_user),
):
    snapshot_id = str(uuid4())
    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                '''
                INSERT INTO generated_cv_snapshots (
                    id,user_id,cv_id,title,target_role,template_key,content
                )
                VALUES (%s,%s,%s,%s,%s,%s,%s::jsonb)
                ''',
                (
                    snapshot_id,
                    user["id"],
                    None,
                    payload.title,
                    payload.target_role,
                    payload.template_key,
                    json.dumps(payload.content, default=str),
                ),
            )
        connection.commit()

    return GeneratedCVResponse(
        title=payload.title,
        target_role=payload.target_role,
        template_key=payload.template_key,
        content=payload.content,
        warnings=[
            "This CV was created from imported information.",
            "Review every fact, date and achievement before submitting it.",
        ],
        snapshot_id=snapshot_id,
    )


@router.post("/generate", response_model=GeneratedCVResponse)
def generate_cv(payload: GenerateCVRequest, user=Depends(get_current_user)):
    content = _build_structured_content(user)
    title = f"{user['full_name']} - {payload.target_role} CV"

    snapshot_id = None
    if payload.save_snapshot:
        snapshot_id = str(uuid4())
        with get_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(
                    '''
                    INSERT INTO generated_cv_snapshots (
                        id,user_id,cv_id,title,target_role,template_key,content
                    )
                    VALUES (%s,%s,%s,%s,%s,%s,%s::jsonb)
                    ''',
                    (
                        snapshot_id, user["id"], payload.cv_id, title,
                        payload.target_role, payload.template_key,
                        json.dumps(content, default=str),
                    ),
                )
            connection.commit()

    return GeneratedCVResponse(
        title=title,
        target_role=payload.target_role,
        template_key=payload.template_key,
        content=content,
        warnings=[
            "Review all generated content before publishing.",
            "The CV was assembled only from saved user records.",
        ],
        snapshot_id=snapshot_id,
    )


@router.post("/ats-score", response_model=ATSScoreResponse)
def ats_score(payload: ATSScoreRequest, user=Depends(get_current_user)):
    result = calculate_ats(payload.cv_content, payload.job_description)

    if payload.save_history:
        with get_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(
                    '''
                    INSERT INTO ats_assessments (
                        id,user_id,cv_id,target_role,job_description,score,
                        matched_keywords,missing_keywords,recommendations
                    )
                    VALUES (%s,%s,%s,%s,%s,%s,%s::jsonb,%s::jsonb,%s::jsonb)
                    ''',
                    (
                        str(uuid4()), user["id"], payload.cv_id, payload.target_role,
                        payload.job_description, result["score"],
                        json.dumps(result["matched_keywords"]),
                        json.dumps(result["missing_keywords"]),
                        json.dumps(result["recommendations"]),
                    ),
                )
            connection.commit()

    return ATSScoreResponse(**result)


@router.get("/ats-history")
def ats_history(
    limit: int = Query(default=20, ge=1, le=100),
    user=Depends(get_current_user),
):
    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                '''
                SELECT *
                FROM ats_assessments
                WHERE user_id=%s
                ORDER BY created_at DESC
                LIMIT %s
                ''',
                (user["id"], limit),
            )
            return cursor.fetchall()


@router.post("/improve-summary", response_model=TextImprovementResponse)
def improve_summary(payload: ImproveSummaryRequest, user=Depends(get_current_user)):
    text, warnings = improve_summary_text(
        payload.target_role,
        payload.current_summary,
        payload.verified_strengths,
    )

    revision_id = None
    if payload.save_revision:
        revision_id = _save_revision(
            user_id=user["id"],
            cv_id=payload.cv_id,
            revision_type="summary",
            source_text=payload.current_summary,
            generated_text=text,
            target_role=payload.target_role,
            warnings=warnings,
        )

    return TextImprovementResponse(
        text=text,
        warnings=warnings,
        revision_id=revision_id,
    )


@router.post("/improve-experience", response_model=TextImprovementResponse)
def improve_experience(payload: ImproveExperienceRequest, user=Depends(get_current_user)):
    text, warnings = improve_experience_text(
        payload.job_title,
        payload.duties,
        payload.verified_achievements,
    )

    revision_id = None
    if payload.save_revision:
        source_text = "\n".join(payload.duties + payload.verified_achievements)
        revision_id = _save_revision(
            user_id=user["id"],
            cv_id=payload.cv_id,
            revision_type="experience",
            source_text=source_text,
            generated_text=text,
            target_role=payload.target_role,
            warnings=warnings,
        )

    return TextImprovementResponse(
        text=text,
        warnings=warnings,
        revision_id=revision_id,
    )


@router.get("/revisions")
def revisions(
    limit: int = Query(default=30, ge=1, le=100),
    user=Depends(get_current_user),
):
    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                '''
                SELECT *
                FROM ai_revisions
                WHERE user_id=%s
                ORDER BY created_at DESC
                LIMIT %s
                ''',
                (user["id"], limit),
            )
            return cursor.fetchall()


@router.put("/revisions/{revision_id}/accept")
def accept_revision(
    revision_id: str,
    payload: RevisionAcceptRequest,
    user=Depends(get_current_user),
):
    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                '''
                UPDATE ai_revisions
                SET accepted=%s
                WHERE id=%s AND user_id=%s
                RETURNING *
                ''',
                (payload.accepted, revision_id, user["id"]),
            )
            row = cursor.fetchone()
        connection.commit()

    if row is None:
        raise HTTPException(status_code=404, detail="Revision not found")
    return row


@router.get("/generated-history")
def generated_history(
    limit: int = Query(default=20, ge=1, le=100),
    user=Depends(get_current_user),
):
    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                '''
                SELECT *
                FROM generated_cv_snapshots
                WHERE user_id=%s
                ORDER BY created_at DESC
                LIMIT %s
                ''',
                (user["id"], limit),
            )
            return cursor.fetchall()


@router.post("/export/docx")
def download_docx(payload: ExportCVRequest, _user=Depends(get_current_user)):
    filename = _safe_filename(payload.filename) + ".docx"
    content = export_docx(payload.cv_content, payload.template_key)

    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/export/pdf")
def download_pdf(payload: ExportCVRequest, _user=Depends(get_current_user)):
    filename = _safe_filename(payload.filename) + ".pdf"
    content = export_pdf(payload.cv_content, payload.template_key)

    return Response(
        content=content,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _build_structured_content(user: dict) -> dict:
    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute("SELECT * FROM profiles WHERE user_id=%s", (user["id"],))
            profile = cursor.fetchone() or {}

            cursor.execute(
                "SELECT * FROM education WHERE user_id=%s ORDER BY start_date DESC NULLS LAST",
                (user["id"],),
            )
            education = cursor.fetchall()

            cursor.execute(
                "SELECT * FROM experience WHERE user_id=%s ORDER BY start_date DESC NULLS LAST",
                (user["id"],),
            )
            experience = cursor.fetchall()

            cursor.execute(
                "SELECT * FROM skills WHERE user_id=%s ORDER BY name",
                (user["id"],),
            )
            skills = cursor.fetchall()

            cursor.execute(
                "SELECT * FROM projects WHERE user_id=%s ORDER BY start_date DESC NULLS LAST",
                (user["id"],),
            )
            projects = cursor.fetchall()

            cursor.execute(
                "SELECT * FROM certifications WHERE user_id=%s ORDER BY issue_date DESC NULLS LAST",
                (user["id"],),
            )
            certifications = cursor.fetchall()

            cursor.execute(
                "SELECT * FROM languages WHERE user_id=%s ORDER BY name",
                (user["id"],),
            )
            languages = cursor.fetchall()

    return {
        "personal_details": {
            "full_name": user["full_name"],
            "email": user["email"],
            "phone": profile.get("phone"),
            "location": profile.get("location"),
            "linkedin_url": profile.get("linkedin_url"),
            "portfolio_url": profile.get("portfolio_url"),
            "website_url": profile.get("website_url"),
        },
        "professional_title": profile.get("professional_title"),
        "professional_summary": profile.get("professional_summary"),
        "skills": skills,
        "experience": experience,
        "education": education,
        "projects": projects,
        "certifications": certifications,
        "languages": languages,
        "references": "Available upon request",
    }


def _save_revision(
    *,
    user_id: str,
    cv_id: str | None,
    revision_type: str,
    source_text: str,
    generated_text: str,
    target_role: str | None,
    warnings: list[str],
) -> str:
    revision_id = str(uuid4())

    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                '''
                INSERT INTO ai_revisions (
                    id,user_id,cv_id,revision_type,source_text,
                    generated_text,target_role,warnings
                )
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s::jsonb)
                ''',
                (
                    revision_id, user_id, cv_id, revision_type,
                    source_text, generated_text, target_role,
                    json.dumps(warnings),
                ),
            )
        connection.commit()

    return revision_id


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "-", value.strip()).strip("-")
    if not cleaned:
        raise HTTPException(status_code=422, detail="Invalid filename")
    return cleaned[:100]
