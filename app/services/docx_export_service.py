from __future__ import annotations

from io import BytesIO

from docx import Document

from app.services.template_renderer import TemplateRenderer


class DOCXExportService:
    """
    Export a CV Studio draft to DOCX.
    """

    def __init__(self):
        self.renderer = TemplateRenderer()

    def export(self, draft: dict) -> bytes:
        data = self.renderer.build_docx(draft)

        document = Document()

        metadata = data["metadata"]

        document.add_heading(metadata["title"], level=1)

        summary = data.get("summary")

        if summary:
            document.add_heading("Professional Summary", level=2)
            document.add_paragraph(summary)

        if data["experience"]:
            document.add_heading("Experience", level=2)

            for item in data["experience"]:
                p = document.add_paragraph()

                p.add_run(item.get("job_title", "")).bold = True

                company = item.get("company", "")

                if company:
                    p.add_run(f" - {company}")

                duration = item.get("duration", "")

                if duration:
                    p.add_run(f" ({duration})")

                for duty in item.get("responsibilities", []):
                    document.add_paragraph(
                        duty,
                        style="List Bullet",
                    )

        if data["education"]:
            document.add_heading("Education", level=2)

            for item in data["education"]:
                document.add_paragraph(
                    f"{item.get('qualification','')} - {item.get('institution','')}"
                )

        if data["skills"]:
            document.add_heading("Skills", level=2)

            for skill in data["skills"]:
                document.add_paragraph(
                    skill,
                    style="List Bullet",
                )

        stream = BytesIO()

        document.save(stream)

        stream.seek(0)

        return stream.read()