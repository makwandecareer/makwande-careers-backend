from __future__ import annotations

import html
import io
import re
from collections import Counter
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer

from app.services.cv_content_normalizer import normalise_cv_content
from app.services.resume_phase14_engine import enrich_phase14

STOPWORDS = {
    "and", "the", "with", "for", "that", "this", "from", "are", "was", "were",
    "have", "has", "will", "your", "you", "our", "their", "they", "into", "using",
    "a", "an", "of", "to", "in", "on", "at", "as", "or", "be", "is", "it",
}


def improve_summary_text(
    target_role: str,
    current_summary: str,
    verified_strengths: list[str],
) -> tuple[str, list[str]]:
    strengths = [item.strip() for item in verified_strengths if item.strip()]
    cleaned_summary = " ".join(current_summary.strip().split()).rstrip(".")

    sections = [f"Results-oriented professional targeting {target_role}"]

    if strengths:
        sections.append(
            "with verified strengths in " + ", ".join(strengths[:6])
        )

    if cleaned_summary:
        sections.append(cleaned_summary)

    result = ". ".join(sections).strip() + "."

    return result, [
        "Review and approve all generated wording before publication.",
        "No qualifications, employers, metrics, dates, or achievements were invented.",
    ]


def improve_experience_text(
    job_title: str,
    duties: list[str],
    verified_achievements: list[str],
) -> tuple[str, list[str]]:
    bullets: list[str] = []

    for duty in duties:
        cleaned = _sentence(duty)
        if cleaned:
            bullets.append(f"• {cleaned}")

    for achievement in verified_achievements:
        cleaned = _sentence(achievement)
        if cleaned:
            bullets.append(f"• Verified achievement: {cleaned}")

    return f"{job_title}\n" + "\n".join(bullets), [
        "Only supplied duties and verified achievements were used."
    ]


def calculate_ats(
    cv_content: dict[str, Any],
    job_description: str,
) -> dict[str, Any]:
    cv_content = normalise_cv_content(cv_content)
    cv_content, intelligence = enrich_phase14(cv_content, job_description)
    cv_text = _flatten_text(cv_content)
    job_keywords = _top_keywords(job_description, 40)
    cv_words = set(_normalise_words(cv_text))
    

    matched = [keyword for keyword in job_keywords if keyword in cv_words]
    missing = [keyword for keyword in job_keywords if keyword not in cv_words]

    keyword_score = int((len(matched) / max(len(job_keywords), 1)) * 70)

    section_score = 0
    for field in ("professional_summary", "skills", "experience", "education"):
        if cv_content.get(field):
            section_score += 7

    score = min(keyword_score + section_score, 100)

    recommendations: list[str] = []

    if missing:
        recommendations.append(
            "Add missing keywords only where they accurately describe verified skills or experience."
        )
    if not cv_content.get("professional_summary"):
        recommendations.append("Add a concise professional summary.")
    if not cv_content.get("skills"):
        recommendations.append("Add a focused skills section.")
    if not cv_content.get("experience"):
        recommendations.append(
            "Add relevant employment, project, internship, or volunteer experience."
        )
    if score >= 80:
        recommendations.append("The CV has strong section and keyword coverage.")

    return {
        "score": score,
        "matched_keywords": matched,
        "missing_keywords": missing[:25],
        "recommendations": recommendations,
        "resume_recommendations": intelligence.get("warnings", []),
        "profile_score": intelligence.get("profile_readiness", 0),
        "resume_intelligence": intelligence,
    }


def export_docx(cv_content: dict[str, Any], template_key: str) -> bytes:
    cv_content = normalise_cv_content(cv_content)
    cv_content, _ = enrich_phase14(cv_content)

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    personal = cv_content.get("personal_details", {})
    full_name = personal.get("full_name") or "Curriculum Vitae"

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(full_name))
    run.bold = True
    run.font.size = Pt(20 if template_key == "executive" else 18)

    professional_title = cv_content.get("professional_title")
    if professional_title:
        line = document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.add_run(str(professional_title)).italic = True

    contact = " | ".join(
        str(item)
        for item in (
            personal.get("email"),
            personal.get("phone"),
            personal.get("location"),
            personal.get("linkedin_url"),
        )
        if item
    )
    if contact:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(contact)

    _docx_section(document, "Professional Summary", cv_content.get("professional_summary"))

    _docx_skill_categories(document, cv_content)

    _docx_experience(document, cv_content.get("experience") or [], template_key)
    _docx_education(document, cv_content.get("education") or [])

    projects = cv_content.get("projects") or []
    if projects:
        document.add_heading("Projects", level=1)
        for item in projects:
            document.add_heading(str(item.get("name") or "Project"), level=2)
            if item.get("description"):
                document.add_paragraph(str(item["description"]))

    certifications = cv_content.get("certifications") or []
    if certifications:
        document.add_heading("Certifications", level=1)
        for item in certifications:
            text = " — ".join(
                part for part in (item.get("name"), item.get("issuer")) if part
            )
            document.add_paragraph(text, style="List Bullet")

    memberships = cv_content.get("professional_memberships") or []
    if memberships:
        document.add_heading("Professional Memberships and Registrations", level=1)
        for item in memberships:
            document.add_paragraph(_professional_membership_text(item), style="List Bullet")

    awards = cv_content.get("awards") or []
    if awards:
        document.add_heading("Awards and Recognition", level=1)
        for item in awards:
            document.add_paragraph(_section_text(item), style="List Bullet")

    volunteer = cv_content.get("volunteer_experience") or []
    if volunteer:
        document.add_heading("Volunteer Experience", level=1)
        for item in volunteer:
            document.add_paragraph(_section_text(item), style="List Bullet")

    languages = cv_content.get("languages") or []
    if languages:
        document.add_heading("Languages", level=1)
        document.add_paragraph(
            " • ".join(
                item.get("name", "") if isinstance(item, dict) else str(item)
                for item in languages
            )
        )

    if cv_content.get("references"):
        _docx_section(
            document,
            "References",
            _section_text(cv_content["references"]),
        )

    if cv_content.get("declaration"):
        _docx_section(
            document,
            "Declaration and Consent",
            cv_content["declaration"],
        )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_pdf(cv_content: dict[str, Any], template_key: str) -> bytes:
    cv_content = normalise_cv_content(cv_content)
    cv_content, _ = enrich_phase14(cv_content)

    buffer = io.BytesIO()
    styles = getSampleStyleSheet()

    body = ParagraphStyle(
        "CVBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        spaceAfter=4,
    )
    heading = ParagraphStyle(
        "CVHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=14,
        spaceBefore=7,
        spaceAfter=4,
    )

    story: list[Any] = []
    personal = cv_content.get("personal_details", {})
    full_name = personal.get("full_name") or "Curriculum Vitae"

    story.append(Paragraph(html.escape(str(full_name)), styles["Title"]))

    professional_title = cv_content.get("professional_title")
    if professional_title:
        story.append(Paragraph(html.escape(str(professional_title)), styles["Heading2"]))

    contact = " | ".join(
        str(item)
        for item in (
            personal.get("email"),
            personal.get("phone"),
            personal.get("location"),
            personal.get("linkedin_url"),
        )
        if item
    )
    if contact:
        story.append(Paragraph(html.escape(contact), body))
        story.append(Spacer(1, 3 * mm))

    _pdf_section(story, heading, body, "Professional Summary", cv_content.get("professional_summary"))

    _pdf_skill_categories(story, heading, body, cv_content)

    experience = cv_content.get("experience") or []
    if experience:
        story.append(Paragraph("Professional Experience", heading))
        for item in experience:
            label = " — ".join(
                str(part) for part in (item.get("job_title"), item.get("company")) if part
            )
            block: list[Any] = [Paragraph(html.escape(label or "Experience"), styles["Heading2"])]
            dates = _date_range(item)
            meta = " | ".join(str(part) for part in (dates, item.get("location"), item.get("employment_type")) if part)
            if meta:
                block.append(Paragraph(html.escape(meta), body))
            duties = item.get("responsibilities") or item.get("duties") or []
            if duties:
                block.append(Paragraph("<b>Key Responsibilities</b>", body))
                for duty in duties:
                    block.append(Paragraph("• " + html.escape(str(duty)), body))
            achievements = item.get("achievements") or []
            if achievements:
                block.append(Paragraph("<b>Key Achievements</b>", body))
                for achievement in achievements:
                    block.append(Paragraph("• " + html.escape(str(achievement)), body))
            story.append(KeepTogether(block))

    education = cv_content.get("education") or []
    if education:
        story.append(Paragraph("Education", heading))
        for item in education:
            label = " — ".join(
                part for part in (item.get("qualification"), item.get("institution")) if part
            )
            story.append(Paragraph(html.escape(label or "Education"), styles["Heading2"]))
            if item.get("field_of_study"):
                story.append(Paragraph(html.escape(str(item["field_of_study"])), body))

    projects = cv_content.get("projects") or []
    if projects:
        story.append(Paragraph("Projects", heading))
        for item in projects:
            story.append(Paragraph(html.escape(str(item.get("name") or "Project")), styles["Heading2"]))
            if item.get("description"):
                story.append(Paragraph(html.escape(str(item["description"])), body))

    certifications = cv_content.get("certifications") or []
    if certifications:
        story.append(Paragraph("Certifications", heading))
        for item in certifications:
            label = " — ".join(
                part for part in (item.get("name"), item.get("issuer")) if part
            )
            story.append(Paragraph("• " + html.escape(label), body))

    memberships = cv_content.get("professional_memberships") or []
    if memberships:
        story.append(Paragraph("Professional Memberships and Registrations", heading))
        for item in memberships:
            story.append(Paragraph("• " + html.escape(_professional_membership_text(item)), body))

    awards = cv_content.get("awards") or []
    if awards:
        story.append(Paragraph("Awards and Recognition", heading))
        for item in awards:
            story.append(Paragraph("• " + html.escape(_section_text(item)), body))

    volunteer = cv_content.get("volunteer_experience") or []
    if volunteer:
        story.append(Paragraph("Volunteer Experience", heading))
        for item in volunteer:
            story.append(Paragraph("• " + html.escape(_section_text(item)), body))

    languages = cv_content.get("languages") or []
    if languages:
        labels = [
            item.get("name", "") if isinstance(item, dict) else str(item)
            for item in languages
        ]
        _pdf_section(story, heading, body, "Languages", " • ".join(filter(None, labels)))

    if cv_content.get("references"):
        _pdf_section(
            story,
            heading,
            body,
            "References",
            _section_text(cv_content["references"]),
        )

    if cv_content.get("declaration"):
        _pdf_section(
            story,
            heading,
            body,
            "Declaration and Consent",
            cv_content["declaration"],
        )

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=str(full_name),
        author="Makwande Careers",
    )
    document.build(story)
    return buffer.getvalue()


def _normalise_words(text: str) -> list[str]:
    words = re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]{1,}", text.lower())
    return [word for word in words if word not in STOPWORDS and len(word) > 2]


def _top_keywords(text: str, limit: int) -> list[str]:
    return [word for word, _ in Counter(_normalise_words(text)).most_common(limit)]


def _flatten_text(value: Any) -> str:
    if isinstance(value, dict):
        return " ".join(_flatten_text(item) for item in value.values())
    if isinstance(value, list):
        return " ".join(_flatten_text(item) for item in value)
    return "" if value is None else str(value)


def _sentence(value: str) -> str:
    cleaned = " ".join(value.strip().split()).rstrip(".")
    if not cleaned:
        return ""
    return cleaned[0].upper() + cleaned[1:] + "."


def _section_text(value: Any) -> str:
    if isinstance(value, list):
        labels: list[str] = []
        for item in value:
            if isinstance(item, dict):
                name = item.get("name") or item.get("full_name")
                relationship = item.get("relationship")
                contact = item.get("contact") or item.get("email") or item.get("phone")
                label = " — ".join(
                    str(part)
                    for part in (name, relationship, contact)
                    if part
                )
                if label:
                    labels.append(label)
            elif item:
                labels.append(str(item))
        return "\n".join(labels)

    if isinstance(value, dict):
        return "\n".join(
            f"{key.replace('_', ' ').title()}: {item}"
            for key, item in value.items()
            if item
        )

    return "" if value is None else str(value)


def _docx_section(document: Document, title: str, value: Any) -> None:
    if value:
        document.add_heading(title, level=1)
        document.add_paragraph(str(value))


def _docx_experience(document: Document, items: list[dict], template_key: str) -> None:
    if not items:
        return
    document.add_heading("Professional Experience", level=1)
    for item in items:
        label = " — ".join(
            str(part) for part in (item.get("job_title"), item.get("company")) if part
        )
        document.add_heading(label or "Experience", level=2)
        meta = " | ".join(
            str(part) for part in (_date_range(item), item.get("location"), item.get("employment_type")) if part
        )
        if meta:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(meta)
            run.italic = True

        duties = item.get("responsibilities") or item.get("duties") or []
        if duties:
            paragraph = document.add_paragraph()
            paragraph.add_run("Key Responsibilities").bold = True
            for duty in duties:
                document.add_paragraph(str(duty), style="List Bullet")

        achievements = item.get("achievements") or []
        if achievements:
            paragraph = document.add_paragraph()
            paragraph.add_run("Key Achievements").bold = True
            for achievement in achievements:
                document.add_paragraph(str(achievement), style="List Bullet")


def _docx_education(document: Document, items: list[dict]) -> None:
    if not items:
        return
    document.add_heading("Education", level=1)
    for item in items:
        label = " — ".join(
            part for part in (item.get("qualification"), item.get("institution")) if part
        )
        document.add_heading(label or "Education", level=2)
        if item.get("field_of_study"):
            document.add_paragraph(str(item["field_of_study"]))


def _date_range(item: dict[str, Any]) -> str:
    start = item.get("start_date") or item.get("date_from")
    end = item.get("end_date") or item.get("date_to")
    if item.get("is_current") and not end:
        end = "Present"
    if start and end:
        return f"({start}) – ({end})"
    return str(start or end or "")


def _professional_membership_text(item: Any) -> str:
    if not isinstance(item, dict):
        return str(item)
    return " | ".join(
        str(value) for value in (
            item.get("name") or item.get("organisation"),
            item.get("membership_number") or item.get("registration_number"),
            item.get("status"),
            item.get("expiry_date"),
        ) if value
    )


def _docx_skill_categories(document: Document, content: dict[str, Any]) -> None:
    categories = content.get("skill_categories") or {}
    labels = {
        "technical_skills": "Technical Skills",
        "software_and_systems": "Software and Systems",
        "laboratory_and_equipment": "Laboratory and Equipment",
        "industry_standards": "Industry Standards",
        "professional_competencies": "Professional Competencies",
    }
    rendered = False
    for key, title in labels.items():
        values = categories.get(key) or []
        if values:
            document.add_heading(title, level=1)
            document.add_paragraph(" • ".join(str(value) for value in values))
            rendered = True
    if not rendered:
        skills = content.get("skills") or []
        names = [item.get("name", "") if isinstance(item, dict) else str(item) for item in skills]
        if any(names):
            document.add_heading("Core Skills", level=1)
            document.add_paragraph(" • ".join(filter(None, names)))


def _pdf_skill_categories(story: list[Any], heading: ParagraphStyle, body: ParagraphStyle, content: dict[str, Any]) -> None:
    categories = content.get("skill_categories") or {}
    labels = {
        "technical_skills": "Technical Skills",
        "software_and_systems": "Software and Systems",
        "laboratory_and_equipment": "Laboratory and Equipment",
        "industry_standards": "Industry Standards",
        "professional_competencies": "Professional Competencies",
    }
    rendered = False
    for key, title in labels.items():
        values = categories.get(key) or []
        if values:
            _pdf_section(story, heading, body, title, " • ".join(str(value) for value in values))
            rendered = True
    if not rendered:
        skills = content.get("skills") or []
        names = [item.get("name", "") if isinstance(item, dict) else str(item) for item in skills]
        if any(names):
            _pdf_section(story, heading, body, "Core Skills", " • ".join(filter(None, names)))


def _pdf_section(
    story: list[Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
    title: str,
    value: Any,
) -> None:
    if value:
        story.append(Paragraph(html.escape(title), heading_style))
        story.append(Paragraph(html.escape(str(value)), body_style))
