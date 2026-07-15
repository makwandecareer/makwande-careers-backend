from __future__ import annotations
import io, re
from collections import Counter
from typing import Any
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from app.schemas_ai_cv import ATSScoreResponse, GeneratedCVResponse, TextImprovementResponse

STOPWORDS={"and","the","with","for","that","this","from","are","was","were","have","has","will","your","you","our","their","they","into","using","a","an","of","to","in","on","at","as","or","be","is","it"}

def _words(text:str)->list[str]:
    return [w for w in re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]{1,}", text.lower()) if w not in STOPWORDS and len(w)>2]

def _top_keywords(text:str, limit:int=30)->list[str]:
    return [w for w,_ in Counter(_words(text)).most_common(limit)]

def improve_summary(*, target_role:str, current_summary:str, verified_strengths:list[str])->TextImprovementResponse:
    strengths=", ".join(x.strip() for x in verified_strengths if x.strip())
    parts=[f"Results-oriented professional targeting {target_role}"]
    if strengths: parts.append(f"with verified strengths in {strengths}")
    if current_summary.strip(): parts.append(" ".join(current_summary.split()).rstrip("."))
    return TextImprovementResponse(text=". ".join(parts)+".", warnings=["Review before publishing.","No facts were invented."])

def improve_experience(*, job_title:str, duties:list[str], verified_achievements:list[str])->TextImprovementResponse:
    lines=[]
    for item in duties:
        x=" ".join(item.split()).rstrip(".")
        if x: lines.append(f"• {x[0].upper()+x[1:]}.")
    for item in verified_achievements:
        x=" ".join(item.split()).rstrip(".")
        if x: lines.append(f"• Achievement: {x[0].upper()+x[1:]}.")
    return TextImprovementResponse(text=f"{job_title}\n"+"\n".join(lines), warnings=["Only supplied information was used."])

def _flatten(value:Any)->str:
    if isinstance(value,dict): return " ".join(_flatten(v) for v in value.values())
    if isinstance(value,list): return " ".join(_flatten(v) for v in value)
    return "" if value is None else str(value)

def calculate_ats_score(cv_content:dict[str,Any], job_description:str)->ATSScoreResponse:
    keywords=_top_keywords(job_description,35)
    cv=set(_words(_flatten(cv_content)))
    matched=[k for k in keywords if k in cv]
    missing=[k for k in keywords if k not in cv]
    score=int(len(matched)/max(len(keywords),1)*70)
    score+=sum(7 for key in ("professional_summary","skills","experience","education") if cv_content.get(key))
    rec=[]
    if missing: rec.append("Add missing keywords only where they truthfully match your background.")
    if not cv_content.get("professional_summary"): rec.append("Add a targeted professional summary.")
    if not cv_content.get("skills"): rec.append("Add a focused skills section.")
    return ATSScoreResponse(score=min(score,100), matched_keywords=matched, missing_keywords=missing[:20], recommendations=rec)

def _serialise(row:dict[str,Any])->dict[str,Any]:
    return {k:(v.isoformat() if hasattr(v,"isoformat") else v) for k,v in row.items()}

def build_generated_cv(*, user:dict[str,Any], profile:dict[str,Any]|None, education:list[dict[str,Any]], experience:list[dict[str,Any]], skills:list[dict[str,Any]], target_role:str, template_key:str)->GeneratedCVResponse:
    profile=profile or {}
    summary=improve_summary(target_role=target_role,current_summary=profile.get("professional_summary") or "",verified_strengths=[x["name"] for x in skills[:8]])
    content={
        "personal_details":{"full_name":user["full_name"],"email":user["email"],"phone":profile.get("phone"),"location":profile.get("location"),"linkedin_url":profile.get("linkedin_url"),"portfolio_url":profile.get("portfolio_url"),"website_url":profile.get("website_url")},
        "professional_title":profile.get("professional_title") or target_role,
        "professional_summary":summary.text,
        "skills":[{"name":x["name"],"proficiency":x.get("proficiency"),"years_experience":float(x["years_experience"]) if x.get("years_experience") is not None else None} for x in skills],
        "experience":[_serialise(x) for x in experience],
        "education":[_serialise(x) for x in education],
        "references":"Available upon request",
    }
    return GeneratedCVResponse(title=f"{user['full_name']} - {target_role} CV",target_role=target_role,template_key=template_key,content=content,warnings=summary.warnings)

def export_docx(cv:dict[str,Any])->bytes:
    doc=Document(); doc.styles["Normal"].font.name="Arial"; doc.styles["Normal"].font.size=Pt(10)
    p=cv.get("personal_details",{}); doc.add_heading(p.get("full_name","Curriculum Vitae"),0)
    contact=" | ".join(str(v) for v in (p.get("email"),p.get("phone"),p.get("location")) if v)
    if contact: doc.add_paragraph(contact)
    for heading,key in (("Professional Summary","professional_summary"),("References","references")):
        if cv.get(key): doc.add_heading(heading,1); doc.add_paragraph(str(cv[key]))
    if cv.get("skills"):
        doc.add_heading("Core Skills",1); doc.add_paragraph(" • ".join((x.get("name","") if isinstance(x,dict) else str(x)) for x in cv["skills"]))
    if cv.get("experience"):
        doc.add_heading("Professional Experience",1)
        for x in cv["experience"]:
            doc.add_heading(" — ".join(v for v in (x.get("job_title"),x.get("company")) if v) or "Experience",2)
            if x.get("description"): doc.add_paragraph(str(x["description"]))
            for a in x.get("achievements") or []: doc.add_paragraph(str(a),style="List Bullet")
    if cv.get("education"):
        doc.add_heading("Education",1)
        for x in cv["education"]: doc.add_heading(" — ".join(v for v in (x.get("qualification"),x.get("institution")) if v) or "Education",2)
    b=io.BytesIO(); doc.save(b); return b.getvalue()

def export_pdf(cv:dict[str,Any])->bytes:
    b=io.BytesIO(); styles=getSampleStyleSheet(); story=[]; p=cv.get("personal_details",{})
    story.append(Paragraph(str(p.get("full_name","Curriculum Vitae")),styles["Title"]))
    contact=" | ".join(str(v) for v in (p.get("email"),p.get("phone"),p.get("location")) if v)
    if contact: story += [Paragraph(contact,styles["BodyText"]),Spacer(1,4*mm)]
    for heading,key in (("Professional Summary","professional_summary"),("References","references")):
        if cv.get(key): story += [Paragraph(heading,styles["Heading1"]),Paragraph(str(cv[key]),styles["BodyText"]),Spacer(1,3*mm)]
    if cv.get("skills"):
        names=" • ".join((x.get("name","") if isinstance(x,dict) else str(x)) for x in cv["skills"])
        story += [Paragraph("Core Skills",styles["Heading1"]),Paragraph(names,styles["BodyText"])]
    if cv.get("experience"):
        story.append(Paragraph("Professional Experience",styles["Heading1"]))
        for x in cv["experience"]:
            story.append(Paragraph(" — ".join(v for v in (x.get("job_title"),x.get("company")) if v) or "Experience",styles["Heading2"]))
            if x.get("description"): story.append(Paragraph(str(x["description"]),styles["BodyText"]))
    if cv.get("education"):
        story.append(Paragraph("Education",styles["Heading1"]))
        for x in cv["education"]: story.append(Paragraph(" — ".join(v for v in (x.get("qualification"),x.get("institution")) if v) or "Education",styles["Heading2"]))
    SimpleDocTemplate(b,pagesize=A4,rightMargin=18*mm,leftMargin=18*mm,topMargin=15*mm,bottomMargin=15*mm).build(story)
    return b.getvalue()
