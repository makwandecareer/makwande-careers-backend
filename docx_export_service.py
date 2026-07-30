from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt

from app.services.template_renderer import TemplateRenderer


class DOCXExportService:
    """
    Production DOCX export service.
    Uses the same render model as the live preview and PDF export.
    """

    def __init__(self):
        self.renderer = TemplateRenderer()

    def export(self, draft) -> bytes:
        model = self.renderer.build_docx(draft)

        document = Document()

        title = document.add_heading(model["metadata"]["title"], level=0)
        title.runs[0].font.size = Pt(22)

        if model["summary"]:
            document.add_heading("Professional Summary", level=1)
            document.add_paragraph(model["summary"])

        if model["experience"]:
            document.add_heading("Professional Experience", level=1)
            for job in model["experience"]:
                p = document.add_paragraph()
                p.add_run(job.get("position", "")).bold = True
                company = job.get("company", "")
                if company:
                    p.add_run(f" | {company}")
                if job.get("duration"):
                    p.add_run(f" ({job['duration']})")
                if job.get("description"):
                    document.add_paragraph(job["description"], style="List Bullet")

        if model["education"]:
            document.add_heading("Education", level=1)
            for edu in model["education"]:
                document.add_paragraph(
                    f"{edu.get('qualification','')} - {edu.get('institution','')}",
                    style="List Bullet",
                )

        if model["skills"]:
            document.add_heading("Skills", level=1)
            skills = []
            for s in model["skills"]:
                if isinstance(s, dict):
                    skills.append(s.get("name", ""))
                else:
                    skills.append(str(s))
            document.add_paragraph(", ".join(filter(None, skills)))

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
